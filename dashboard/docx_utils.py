from docx import Document
from docx.oxml.ns import qn
from docx.shared import RGBColor

def extract_editable_runs(docx_path):
    """
    Extracts all runs from the DOCX, marking those that are underlined or highlighted as editable.
    Returns a list of paragraphs, each as a list of dicts: {text, editable, run_idx}.
    """
    doc = Document(docx_path)
    result = []
    for para in doc.paragraphs:
        para_runs = []
        for idx, run in enumerate(para.runs):
            is_underlined = run.underline is True
            is_highlighted = False
            if run._element.rPr is not None:
                highlight = run._element.rPr.find(qn('w:highlight'))
                if highlight is not None:
                    is_highlighted = True
            para_runs.append({
                'text': run.text,
                'editable': is_underlined or is_highlighted,
                'run_idx': idx
            })
        result.append(para_runs)
    return result

def update_editable_runs(docx_path, updates, output_path):
    """
    Updates the DOCX file's underlined/highlighted runs with new text from updates dict {(para_idx, run_idx): new_text}.
    """
    doc = Document(docx_path)
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            key = (para_idx, run_idx)
            if key in updates:
                run.text = updates[key]
    doc.save(output_path)
