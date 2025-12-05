from docx import Document
from docx.oxml.ns import qn

def extract_editable_runs_by_page(docx_path):
    """
    Extracts all runs from the DOCX, splitting by page breaks.
    Returns a list of pages, each page is a list of paragraphs, each as a list of dicts: {text, editable, run_idx, para_idx}.
    """
    doc = Document(docx_path)
    pages = []
    current_page = []
    para_idx = 0
    for para in doc.paragraphs:
        para_runs = []
        for idx, run in enumerate(para.runs):
            is_underlined = run.underline is True
            is_highlighted = False
            if run._element.rPr is not None:
                highlight = run._element.rPr.find(qn('w:highlight'))
                if highlight is not None:
                    is_highlighted = True
            para_runs.append({
                'text': run.text,
                'editable': is_underlined or is_highlighted,
                'run_idx': idx,
                'para_idx': para_idx
            })
            # Check for page break in this run
            for br in run._element.findall('.//w:br', namespaces=run._element.nsmap):
                if br.get(qn('w:type')) == 'page':
                    current_page.append(para_runs)
                    pages.append(current_page)
                    current_page = []
                    para_runs = []
        if para_runs:
            current_page.append(para_runs)
        para_idx += 1
    if current_page:
        pages.append(current_page)
    return pages

def update_editable_runs_by_page(docx_path, updates, output_path):
    """
    Updates the DOCX file's underlined/highlighted runs with new text from updates dict {(para_idx, run_idx): new_text}.
    """
    doc = Document(docx_path)
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            key = (para_idx, run_idx)
            if key in updates:
                run.text = updates[key]
    doc.save(output_path)
