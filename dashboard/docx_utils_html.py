import io
from docx import Document
import mammoth
from docx.oxml.ns import qn

def extract_docx_page_html(docx_path, page_idx, markers):
    """
    Extracts a single page (by markers) from a DOCX and converts it to HTML (tables and paragraphs preserved).
    Can be reused for any document type.
    """
    doc = Document(docx_path)
    body = doc.element.body
    blocks = []
    for child in body.iterchildren():
        if child.tag == qn('w:p'):
            blocks.append(('p', child))
        elif child.tag == qn('w:tbl'):
            blocks.append(('tbl', child))
    pages = []
    current_page = []
    marker_idx = 0
    marker_count = len(markers)
    for block_type, block in blocks:
        if block_type == 'p':
            para_text = ''
            for node in block.iter():
                if node.tag == qn('w:t'):
                    para_text += node.text or ''
            current_page.append(('p', block))
            if marker_idx < marker_count and markers[marker_idx] in para_text.strip():
                if current_page:
                    pages.append(current_page)
                current_page = []
                marker_idx += 1
        elif block_type == 'tbl':
            current_page.append(('tbl', block))
    if current_page:
        pages.append(current_page)
    # Clamp page_idx
    if page_idx < 1:
        page_idx = 1
    if page_idx > len(pages):
        page_idx = len(pages)
    # Build a temp DOCX for the selected page, preserving paragraphs and tables
    temp_doc = Document()
    temp_doc._body.clear_content()
    for block_type, block in pages[page_idx-1]:
        temp_doc._body._element.append(block)
    temp_stream = io.BytesIO()
    temp_doc.save(temp_stream)
    temp_stream.seek(0)
    result = mammoth.convert_to_html(temp_stream)
    html = result.value
    return html, len(pages)

def extract_docx_full_html(docx_path):
    """
    Converts the entire DOCX to HTML (best effort, tables/paragraphs preserved).
    Can be reused for any document type.
    """
    try:
        with open(docx_path, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value
            if html.strip():
                return html
    except Exception:
        pass
    doc = Document(docx_path)
    return '<br>'.join([p.text for p in doc.paragraphs])
