"""
AI Resume Parser — Extracts text and keywords from student CV files (PDF/DOCX).
Used by the matching engine to compute AI-powered resume-to-internship similarity.
"""

import io
import re
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_field):
    """Extract raw text from a PDF file field."""
    try:
        from PyPDF2 import PdfReader

        file_field.open("rb")
        reader = PdfReader(io.BytesIO(file_field.read()))
        file_field.close()

        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n".join(text_parts)
    except Exception as e:
        logger.warning("Failed to extract text from PDF: %s", e)
        return ""


def extract_text_from_docx(file_field):
    """Extract raw text from a DOCX file field."""
    try:
        from docx import Document

        file_field.open("rb")
        doc = Document(io.BytesIO(file_field.read()))
        file_field.close()

        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())

        # Also extract from tables (resumes often use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())

        return "\n".join(text_parts)
    except Exception as e:
        logger.warning("Failed to extract text from DOCX: %s", e)
        return ""


def extract_text_from_cv(file_field):
    """
    Extract raw text from a student's CV file (PDF or DOCX).
    Returns empty string if extraction fails or file type is unsupported.
    """
    if not file_field:
        return ""

    name = (file_field.name or "").lower()

    if name.endswith(".pdf"):
        return extract_text_from_pdf(file_field)
    elif name.endswith(".docx"):
        return extract_text_from_docx(file_field)
    elif name.endswith(".doc"):
        # Legacy .doc files can't be easily parsed without LibreOffice
        logger.info("Legacy .doc file detected — skipping text extraction: %s", name)
        return ""
    else:
        return ""


def clean_text(text):
    """Normalize and clean extracted text for NLP processing."""
    if not text:
        return ""

    # Lowercase
    text = text.lower()

    # Remove emails, URLs, phone numbers
    text = re.sub(r"\S+@\S+\.\S+", " ", text)
    text = re.sub(r"https?://\S+", " ", text)
    text = re.sub(r"\+?\d[\d\s\-()]{7,}", " ", text)

    # Remove special characters but keep spaces and alphanumeric
    text = re.sub(r"[^a-z0-9\s]", " ", text)

    # Collapse whitespace
    text = re.sub(r"\s+", " ", text).strip()

    return text


# Common stop words to filter out from keyword extraction
STOP_WORDS = frozenset({
    "a", "an", "the", "and", "or", "but", "in", "on", "at", "to", "for",
    "of", "with", "by", "from", "is", "are", "was", "were", "be", "been",
    "being", "have", "has", "had", "do", "does", "did", "will", "would",
    "could", "should", "may", "might", "shall", "can", "need", "must",
    "i", "me", "my", "we", "our", "you", "your", "he", "she", "it",
    "they", "them", "their", "this", "that", "these", "those", "am",
    "not", "no", "nor", "so", "if", "up", "out", "about", "into",
    "over", "after", "under", "between", "through", "during", "before",
    "above", "below", "each", "every", "all", "both", "few", "more",
    "most", "other", "some", "such", "than", "too", "very", "just",
    "also", "as", "its", "own", "same", "which", "who", "whom",
    "what", "when", "where", "why", "how", "any", "here", "there",
})


def extract_keywords(text):
    """
    Extract meaningful keywords from cleaned text.
    Returns a set of unique keywords (lowercased, stop words removed).
    """
    cleaned = clean_text(text)
    if not cleaned:
        return set()

    words = cleaned.split()
    # Filter: remove stop words, very short words, and pure numbers
    keywords = {
        w for w in words
        if w not in STOP_WORDS and len(w) > 1 and not w.isdigit()
    }

    return keywords


def parse_student_resume(student_profile):
    """
    Parse a student's uploaded CV and return structured data for matching.

    Returns dict with:
        - raw_text: The full extracted text
        - cleaned_text: Normalized text for NLP
        - keywords: Set of extracted keywords
        - has_resume: Whether resume was successfully parsed
    """
    result = {
        "raw_text": "",
        "cleaned_text": "",
        "keywords": set(),
        "has_resume": False,
    }

    if not student_profile.cv:
        return result

    try:
        raw_text = extract_text_from_cv(student_profile.cv)
        if raw_text:
            result["raw_text"] = raw_text
            result["cleaned_text"] = clean_text(raw_text)
            result["keywords"] = extract_keywords(raw_text)
            result["has_resume"] = True
    except Exception as e:
        logger.error("Error parsing resume for %s: %s", student_profile, e)

    return result


def build_internship_text(internship):
    """
    Build a combined text representation of an internship for TF-IDF comparison.
    Includes: title, description, required skills, company info.
    """
    parts = []

    # Internship title and description (weighted by repeating)
    parts.append(internship.title)
    parts.append(internship.title)  # Double weight for title
    parts.append(internship.description)

    # Required skills (triple weight — most important for matching)
    skill_names = list(internship.required_skills.values_list("name", flat=True))
    for skill_name in skill_names:
        parts.append(skill_name)
        parts.append(skill_name)
        parts.append(skill_name)

    # Company info
    company = internship.company
    parts.append(company.name)
    parts.append(company.description)

    return clean_text(" ".join(parts))
